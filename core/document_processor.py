"""
Document Processing - Handle PDF, DOCX, TXT, and Markdown files
Optimized chunking and text extraction for Apple Silicon
"""

import asyncio
import hashlib
import json
import logging
import os
import threading
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

import aiofiles
from fastapi import UploadFile
from watchdog.events import FileSystemEventHandler
from watchdog.observers import Observer

# Document processing libraries
try:
    import markdown
    import PyPDF2
    from bs4 import BeautifulSoup
    from docx import Document

    PROCESSING_AVAILABLE = True
except ImportError as e:
    PROCESSING_AVAILABLE = False
    print(f"⚠️  Document processing libraries not available: {e}")

from config import settings


class DocumentChangeHandler(FileSystemEventHandler):
    """Handler for file system events in monitored folders"""

    def __init__(self, document_processor):
        self.document_processor = document_processor
        self.logger = logging.getLogger(__name__)
        # Debounce timer to avoid processing rapid file changes
        self._pending_files = {}
        self._timer = None

    def on_created(self, event):
        if not event.is_directory:
            self._schedule_processing(event.src_path, "created")

    def on_modified(self, event):
        if not event.is_directory:
            self._schedule_processing(event.src_path, "modified")

    def on_deleted(self, event):
        if not event.is_directory:
            self._schedule_processing(event.src_path, "deleted")

    def _schedule_processing(self, file_path: str, event_type: str):
        """Schedule file processing with debouncing"""
        self._pending_files[file_path] = {
            "event_type": event_type,
            "timestamp": time.time(),
        }

        # Cancel existing timer
        if self._timer:
            self._timer.cancel()

        # Schedule processing after 2 seconds delay
        self._timer = threading.Timer(2.0, self._process_pending_files)
        self._timer.start()

    def _process_pending_files(self):
        """Process all pending file changes"""
        if not self._pending_files:
            return

        files_to_process = self._pending_files.copy()
        self._pending_files.clear()

        # Process in a separate thread to avoid blocking the file watcher
        threading.Thread(
            target=self._async_process_files, args=(files_to_process,), daemon=True
        ).start()

    def _async_process_files(self, files_to_process):
        """Process files asynchronously"""
        for file_path, info in files_to_process.items():
            try:
                if info["event_type"] == "deleted":
                    asyncio.run(
                        self.document_processor.remove_document_from_store(file_path)
                    )
                else:
                    # Check if file still exists and is a supported format
                    if os.path.exists(
                        file_path
                    ) and self.document_processor.is_supported_file(file_path):
                        asyncio.run(
                            self.document_processor.auto_process_document(file_path)
                        )

            except Exception as e:
                self.logger.error(f"Error auto-processing {file_path}: {e}")


class FileWatcher:
    """File system watcher for monitoring document folders"""

    def __init__(self, document_processor):
        self.document_processor = document_processor
        self.observer = Observer()
        self.monitored_paths = set()
        self.handler = DocumentChangeHandler(document_processor)
        self.logger = logging.getLogger(__name__)
        self.is_running = False

    def add_monitored_path(self, path: str) -> bool:
        """Add a path to be monitored"""
        try:
            path_obj = Path(path).resolve()
            if not path_obj.exists():
                self.logger.warning(f"Path does not exist: {path}")
                return False

            if not path_obj.is_dir():
                self.logger.warning(f"Path is not a directory: {path}")
                return False

            path_str = str(path_obj)
            if path_str not in self.monitored_paths:
                self.observer.schedule(self.handler, path_str, recursive=True)
                self.monitored_paths.add(path_str)
                self.logger.info(f"✅ Now monitoring: {path_str}")

                # Initial scan of the folder
                asyncio.create_task(self._initial_scan(path_str))

            return True

        except Exception as e:
            self.logger.error(f"Error adding monitored path {path}: {e}")
            return False

    def remove_monitored_path(self, path: str) -> bool:
        """Remove a path from monitoring"""
        try:
            path_obj = Path(path).resolve()
            path_str = str(path_obj)

            if path_str in self.monitored_paths:
                # Note: watchdog doesn't have a direct way to unschedule a specific path
                # We'd need to rebuild the observer, but for now we'll just remove from our tracking
                self.monitored_paths.discard(path_str)
                self.logger.info(f"🗑️ Removed from monitoring: {path_str}")
                return True

            return False

        except Exception as e:
            self.logger.error(f"Error removing monitored path {path}: {e}")
            return False

    def start(self):
        """Start the file watcher"""
        if not self.is_running:
            self.observer.start()
            self.is_running = True
            self.logger.info("🔍 File watcher started")

    def stop(self):
        """Stop the file watcher"""
        if self.is_running:
            self.observer.stop()
            self.observer.join()
            self.is_running = False
            self.logger.info("⏹️ File watcher stopped")

    def get_monitored_paths(self) -> List[str]:
        """Get list of currently monitored paths"""
        return list(self.monitored_paths)

    async def _initial_scan(self, folder_path: str):
        """Perform initial scan of a newly monitored folder"""
        try:
            self.logger.info(f"🔄 Initial scan of {folder_path}")
            path_obj = Path(folder_path)

            for file_path in path_obj.rglob("*"):
                if file_path.is_file() and self.document_processor.is_supported_file(
                    str(file_path)
                ):
                    await self.document_processor.auto_process_document(str(file_path))

            self.logger.info(f"✅ Initial scan complete for {folder_path}")

        except Exception as e:
            self.logger.error(f"Error during initial scan of {folder_path}: {e}")


class DocumentProcessor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.documents_dir = Path("documents")
        self.documents_dir.mkdir(exist_ok=True)

        # File watcher for automatic monitoring
        self.file_watcher = None
        self.supported_extensions = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}

        # Document tracking
        self.processed_documents = {}
        self.document_metadata = {}

    async def initialize_watcher(self):
        """Initialize the file watcher"""
        if not self.file_watcher:
            self.file_watcher = FileWatcher(self)
            self.file_watcher.start()

    def is_supported_file(self, file_path: str) -> bool:
        """Check if file has supported extension"""
        return Path(file_path).suffix.lower() in self.supported_extensions

    async def auto_process_document(self, file_path: str) -> Dict[str, Any]:
        """Automatically process a document detected by file watcher"""
        try:
            self.logger.info(f"🔄 Auto-processing: {file_path}")

            # Check if file exists and is readable
            path_obj = Path(file_path)
            if not path_obj.exists():
                self.logger.warning(f"File no longer exists: {file_path}")
                return {"status": "error", "message": "File not found"}

            # Extract text from the file
            text_content = await self._extract_text_from_file(str(path_obj))
            if not text_content:
                self.logger.warning(f"No text extracted from: {file_path}")
                return {"status": "error", "message": "No text content"}

            # Process the document
            doc_id = self._generate_document_id(str(path_obj))
            chunks = self._create_chunks(text_content)

            # Get vector store and add chunks
            from core.vector_store import get_vector_store

            vector_store = get_vector_store()

            # Remove existing document first (in case of update)
            await self.remove_document_from_store(file_path)

            # Add document chunks to vector store (use active library)
            chunks_with_metadata = []
            for i, chunk in enumerate(chunks):
                metadata = {
                    "source": str(path_obj),
                    "chunk_id": f"{doc_id}_{i}",
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_name": path_obj.name,
                    "file_size": path_obj.stat().st_size,
                    "last_modified": datetime.fromtimestamp(
                        path_obj.stat().st_mtime
                    ).isoformat(),
                    "auto_processed": True,
                    "text": chunk,
                }
                chunks_with_metadata.append(metadata)

            # Add to active library
            await vector_store.add_documents(chunks_with_metadata)

            self.logger.info(f"✅ Auto-processed {file_path}: {len(chunks)} chunks")

            return {
                "status": "success",
                "file_path": str(path_obj),
                "chunks_created": len(chunks),
                "document_id": doc_id,
                "auto_processed": True,
            }

        except Exception as e:
            self.logger.error(f"Error auto-processing {file_path}: {e}")
            return {"status": "error", "message": str(e)}

    async def remove_document_from_store(self, file_path: str) -> bool:
        """Remove a document and its chunks from the vector store"""
        try:
            from core.vector_store import get_vector_store

            vector_store = get_vector_store()

            # Get the normalized path
            path_obj = Path(file_path).resolve()

            # Find and remove all chunks for this document
            removed_count = vector_store.remove_by_source(str(path_obj))

            if removed_count > 0:
                self.logger.info(f"🗑️ Removed {removed_count} chunks for {file_path}")

            return removed_count > 0

        except Exception as e:
            self.logger.error(f"Error removing document {file_path}: {e}")
            return False

    async def add_monitored_folder(self, folder_path: str) -> Dict[str, Any]:
        """Add a folder to continuous monitoring"""
        try:
            if not self.file_watcher:
                await self.initialize_watcher()

            success = self.file_watcher.add_monitored_path(folder_path)

            if success:
                return {
                    "status": "success",
                    "message": f"Now monitoring: {folder_path}",
                    "monitored_paths": self.file_watcher.get_monitored_paths(),
                }
            else:
                return {
                    "status": "error",
                    "message": f"Failed to monitor: {folder_path}",
                }

        except Exception as e:
            self.logger.error(f"Error adding monitored folder {folder_path}: {e}")
            return {"status": "error", "message": str(e)}

    async def remove_monitored_folder(self, folder_path: str) -> Dict[str, Any]:
        """Remove a folder from monitoring"""
        try:
            if not self.file_watcher:
                return {"status": "error", "message": "File watcher not initialized"}

            success = self.file_watcher.remove_monitored_path(folder_path)

            if success:
                return {
                    "status": "success",
                    "message": f"Stopped monitoring: {folder_path}",
                    "monitored_paths": self.file_watcher.get_monitored_paths(),
                }
            else:
                return {
                    "status": "error",
                    "message": f"Path not found in monitoring: {folder_path}",
                }

        except Exception as e:
            self.logger.error(f"Error removing monitored folder {folder_path}: {e}")
            return {"status": "error", "message": str(e)}

    async def get_monitoring_status(self) -> Dict[str, Any]:
        """Get current monitoring status and paths"""
        try:
            if not self.file_watcher:
                return {
                    "status": "success",
                    "watcher_running": False,
                    "monitored_paths": [],
                    "supported_extensions": list(self.supported_extensions),
                }

            return {
                "status": "success",
                "watcher_running": self.file_watcher.is_running,
                "monitored_paths": self.file_watcher.get_monitored_paths(),
                "supported_extensions": list(self.supported_extensions),
            }

        except Exception as e:
            self.logger.error(f"Error getting monitoring status: {e}")
            return {"status": "error", "message": str(e)}

    async def process_file(self, file: UploadFile) -> Dict[str, Any]:
        """Process uploaded file and extract text chunks"""
        try:
            print(f"📄 Processing file: {file.filename}")

            # Read file content
            content = await file.read()
            file_hash = hashlib.md5(content).hexdigest()

            # Check if already processed
            if file_hash in self.processed_documents:
                print(f"♻️  File already processed: {file.filename}")
                return self.processed_documents[file_hash]

            # Save file
            file_path = settings.documents_dir / f"{file_hash}_{file.filename}"
            async with aiofiles.open(file_path, "wb") as f:
                await f.write(content)

            # Extract text based on file type
            extension = Path(file.filename).suffix.lower()

            if extension == ".pdf":
                text = await self._extract_pdf_text(content)
            elif extension == ".docx":
                text = await self._extract_docx_text(content)
            elif extension in [".txt", ".md"]:
                text = content.decode("utf-8", errors="ignore")
                if extension == ".md":
                    text = await self._process_markdown(text)
            elif extension == ".html":
                text = await self._extract_html_text(
                    content.decode("utf-8", errors="ignore")
                )
            else:
                raise ValueError(f"Unsupported file type: {extension}")

            if not text.strip():
                raise ValueError(f"No text content found in {file.filename}")

            # Create intelligent chunks
            chunks = await self._create_intelligent_chunks(text, file.filename)

            # Store processing result
            result = {
                "filename": file.filename,
                "file_hash": file_hash,
                "file_path": str(file_path),
                "text_length": len(text),
                "chunks": chunks,
                "chunk_count": len(chunks),
                "extension": extension,
            }

            self.processed_documents[file_hash] = result
            print(
                f"✅ Processed {file.filename}: {len(chunks)} chunks, {len(text)} characters"
            )
            return result

        except Exception as e:
            print(f"❌ Error processing file {file.filename}: {e}")
            raise

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            import io

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text + "\n"
                except Exception as e:
                    print(f"⚠️  Error extracting page {page_num + 1}: {e}")
                    continue

            return text.strip()
        except Exception as e:
            print(f"❌ Error extracting PDF text: {e}")
            raise

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            import io

            doc = Document(io.BytesIO(content))
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            return text.strip()
        except Exception as e:
            print(f"❌ Error extracting DOCX text: {e}")
            raise

    async def _process_markdown(self, text: str) -> str:
        """Process Markdown and convert to plain text"""
        try:
            html = markdown.markdown(text)
            soup = BeautifulSoup(html, "html.parser")

            # Extract text while preserving some structure
            plain_text = ""
            for element in soup.find_all(
                ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]
            ):
                if element.name.startswith("h"):
                    plain_text += f"\n## {element.get_text().strip()}\n"
                else:
                    plain_text += element.get_text().strip() + "\n"

            return plain_text.strip()
        except Exception as e:
            print(f"❌ Error processing Markdown: {e}")
            return text

    async def _extract_html_text(self, html: str) -> str:
        """Extract text from HTML"""
        try:
            soup = BeautifulSoup(html, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text with some structure preservation
            text = ""
            for element in soup.find_all(
                ["h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "li"]
            ):
                element_text = element.get_text().strip()
                if element_text:
                    if element.name.startswith("h"):
                        text += f"\n## {element_text}\n"
                    else:
                        text += element_text + "\n"

            return text.strip()
        except Exception as e:
            print(f"❌ Error extracting HTML text: {e}")
            return ""

    async def _create_intelligent_chunks(
        self, text: str, filename: str
    ) -> List[Dict[str, Any]]:
        """Create intelligent text chunks with overlap and context preservation"""
        chunks = []

        # Split text into paragraphs first
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        if not paragraphs:
            # Fallback to line-based splitting
            paragraphs = [line.strip() for line in text.split("\n") if line.strip()]

        current_chunk = ""
        current_chunk_paragraphs = []
        chunk_id = 0

        for i, paragraph in enumerate(paragraphs):
            # Check if adding this paragraph exceeds chunk size
            potential_chunk = (
                current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            )

            if len(potential_chunk) > settings.chunk_size and current_chunk:
                # Save current chunk
                chunk_text = current_chunk.strip()

                chunks.append(
                    {
                        "id": f"{filename}_chunk_{chunk_id}",
                        "text": chunk_text,
                        "source": filename,
                        "chunk_index": chunk_id,
                        "metadata": {
                            "filename": filename,
                            "chunk_size": len(chunk_text),
                            "chunk_index": chunk_id,
                            "paragraph_count": len(current_chunk_paragraphs),
                            "overlap_with_next": False,
                        },
                    }
                )

                # Create overlap with previous chunk
                overlap_size = min(
                    settings.chunk_overlap, len(current_chunk_paragraphs)
                )
                if overlap_size > 0 and len(current_chunk_paragraphs) > 1:
                    # Take last few paragraphs for overlap
                    overlap_paragraphs = current_chunk_paragraphs[-overlap_size:]
                    current_chunk = "\n\n".join(overlap_paragraphs)
                    current_chunk_paragraphs = overlap_paragraphs[:]
                    chunks[-1]["metadata"]["overlap_with_next"] = True
                else:
                    current_chunk = ""
                    current_chunk_paragraphs = []

                chunk_id += 1

            # Add paragraph to current chunk
            if current_chunk:
                current_chunk += "\n\n" + paragraph
            else:
                current_chunk = paragraph
            current_chunk_paragraphs.append(paragraph)

        # Add the last chunk if it exists
        if current_chunk.strip():
            chunks.append(
                {
                    "id": f"{filename}_chunk_{chunk_id}",
                    "text": current_chunk.strip(),
                    "source": filename,
                    "chunk_index": chunk_id,
                    "metadata": {
                        "filename": filename,
                        "chunk_size": len(current_chunk),
                        "chunk_index": chunk_id,
                        "paragraph_count": len(current_chunk_paragraphs),
                        "overlap_with_next": False,
                    },
                }
            )

        # If no chunks were created (very short text), create one chunk
        if not chunks and text.strip():
            chunks.append(
                {
                    "id": f"{filename}_chunk_0",
                    "text": text.strip(),
                    "source": filename,
                    "chunk_index": 0,
                    "metadata": {
                        "filename": filename,
                        "chunk_size": len(text),
                        "chunk_index": 0,
                        "paragraph_count": 1,
                        "overlap_with_next": False,
                    },
                }
            )

        return chunks

    async def get_document_count(self) -> int:
        """Get total number of processed documents"""
        return len(self.processed_documents)

    async def get_document_list(self) -> List[Dict[str, Any]]:
        """Get list of all processed documents"""
        return [
            {
                "filename": doc["filename"],
                "file_hash": doc["file_hash"],
                "text_length": doc["text_length"],
                "chunk_count": doc["chunk_count"],
                "extension": doc.get("extension", "unknown"),
            }
            for doc in self.processed_documents.values()
        ]

    async def delete_document(self, file_hash: str) -> bool:
        """Delete a processed document"""
        try:
            if file_hash in self.processed_documents:
                # Remove file
                file_path = Path(self.processed_documents[file_hash]["file_path"])
                if file_path.exists():
                    file_path.unlink()

                # Remove from memory
                filename = self.processed_documents[file_hash]["filename"]
                del self.processed_documents[file_hash]
                print(f"🗑️  Deleted document: {filename}")
                return True
            return False
        except Exception as e:
            print(f"❌ Error deleting document: {e}")
            return False

    async def get_document_by_hash(self, file_hash: str) -> Dict[str, Any]:
        """Get document details by hash"""
        return self.processed_documents.get(file_hash, {})

    async def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return self.supported_extensions

    async def reprocess_document(self, file_hash: str) -> Dict[str, Any]:
        """Reprocess an existing document with current settings"""
        try:
            if file_hash not in self.processed_documents:
                raise ValueError("Document not found")

            doc_info = self.processed_documents[file_hash]
            file_path = Path(doc_info["file_path"])

            if not file_path.exists():
                raise ValueError("Document file not found on disk")

            # Read file and reprocess
            async with aiofiles.open(file_path, "rb") as f:
                content = await f.read()

            # Create a mock UploadFile object
            class MockUploadFile:
                def __init__(self, content, filename):
                    self.content = content
                    self.filename = filename

                async def read(self):
                    return self.content

            mock_file = MockUploadFile(content, doc_info["filename"])

            # Remove old entry and reprocess
            del self.processed_documents[file_hash]
            result = await self.process_file(mock_file)

            print(f"♻️  Reprocessed document: {doc_info['filename']}")
            return result

        except Exception as e:
            print(f"❌ Error reprocessing document: {e}")
            raise

    async def process_folder_batch(self, files: List[UploadFile]) -> Dict[str, Any]:
        """Process multiple files from folder upload with detailed reporting"""
        try:
            print(f"📁 Starting batch processing of {len(files)} files...")

            successful_results = []
            failed_results = []
            supported_files = 0
            skipped_files = 0

            for i, file in enumerate(files):
                try:
                    print(f"📄 Processing file {i+1}/{len(files)}: {file.filename}")

                    # Check if file type is supported
                    extension = Path(file.filename).suffix.lower()
                    if extension not in self.supported_extensions:
                        skipped_files += 1
                        print(
                            f"⏭️  Skipping unsupported file: {file.filename} ({extension})"
                        )
                        continue

                    supported_files += 1

                    # Process the file
                    result = await self.process_file(file)
                    successful_results.append(
                        {
                            "filename": result["filename"],
                            "file_hash": result["file_hash"],
                            "chunk_count": result["chunk_count"],
                            "text_length": result["text_length"],
                            "extension": result.get("extension", "unknown"),
                            "chunks": result[
                                "chunks"
                            ],  # Include chunks for vector store
                        }
                    )

                    print(
                        f"✅ Successfully processed: {file.filename} ({result['chunk_count']} chunks)"
                    )

                except Exception as e:
                    failed_results.append(
                        {
                            "filename": file.filename,
                            "error": str(e),
                            "extension": Path(file.filename).suffix.lower(),
                        }
                    )
                    print(f"❌ Failed to process: {file.filename} - {str(e)}")
                    continue

            # Compile results
            stats = {
                "total_files": len(files),
                "successful_files": len(successful_results),
                "failed_files": len(failed_results),
                "skipped_files": skipped_files,
                "supported_files": supported_files,
                "total_chunks": sum(r["chunk_count"] for r in successful_results),
                "total_text_length": sum(r["text_length"] for r in successful_results),
            }

            print(f"📊 Batch processing complete:")
            print(f"   ✅ Successful: {stats['successful_files']}")
            print(f"   ❌ Failed: {stats['failed_files']}")
            print(f"   ⏭️  Skipped: {stats['skipped_files']}")
            print(f"   📝 Total chunks: {stats['total_chunks']}")

            return {
                "successful": successful_results,
                "failed": failed_results,
                "stats": stats,
            }

        except Exception as e:
            print(f"❌ Error in batch processing: {e}")
            raise
